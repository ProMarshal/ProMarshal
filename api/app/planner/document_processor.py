"""Document processing utilities for extracting text from various file formats"""

import io
from typing import BinaryIO
import docx
from pypdf import PdfReader


class DocumentProcessor:
    """Process different document formats and extract text"""

    @staticmethod
    def extract_text_from_docx(file_content: BinaryIO) -> str:
        """Extract text from Word document (.docx)"""
        try:
            doc = docx.Document(file_content)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Error extracting text from Word document: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(file_content: BinaryIO) -> str:
        """Extract text from PDF document"""
        try:
            reader = PdfReader(file_content)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_content: BinaryIO) -> str:
        """Extract text from plain text or markdown file"""
        try:
            content = file_content.read()

            # Try UTF-8 first
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                return content.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Error extracting text from text file: {str(e)}")

    @classmethod
    def extract_text(cls, file_content: BinaryIO, filename: str) -> str:
        """
        Extract text from file based on extension

        Args:
            file_content: Binary file content
            filename: Original filename with extension

        Returns:
            Extracted text content
        """
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

        if file_ext in ['docx', 'doc']:
            return cls.extract_text_from_docx(file_content)
        elif file_ext == 'pdf':
            return cls.extract_text_from_pdf(file_content)
        elif file_ext in ['txt', 'md', 'markdown', 'text']:
            return cls.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
